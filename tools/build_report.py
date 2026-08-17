from __future__ import annotations

import csv
import math
import re
from collections import defaultdict
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_MD = ROOT / "REPORT.md"
OUT_DIR = ROOT.parent / "output" / "report"
FIG_DIR = OUT_DIR / "figures"
OUT_DOCX = OUT_DIR / "NBQSS_2026_Tomography_Report.docx"
NAVY = RGBColor(20, 52, 78)
TEAL = RGBColor(15, 113, 115)
GOLD = RGBColor(179, 129, 32)
INK = RGBColor(32, 38, 45)
MUTED = RGBColor(93, 105, 116)
LIGHT = "EAF0F4"
PALE = "F5F7F9"
WHITE = "FFFFFF"


def set_font(run, name="Arial", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = 0
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def create_decimal_numbering(doc):
    """Create a fresh real Word numbering sequence starting at 1."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    p_pr.append(ind)
    level.append(p_pr)
    numbering.append(abstract)
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 18, 10),
        ("Heading 2", 13, TEAL, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208
    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Menlo"
    code.font.size = Pt(8.5)
    code.paragraph_format.left_indent = Inches(0.25)
    code.paragraph_format.right_indent = Inches(0.25)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(3)
    code.paragraph_format.line_spacing = 1.05
    if "Caption" in styles:
        caption = styles["Caption"]
        caption.font.name = "Arial"
        caption.font.size = Pt(9)
        caption.font.italic = True
        caption.font.color.rgb = MUTED
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_before = Pt(4)
        caption.paragraph_format.space_after = Pt(8)


def find_font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def load_summary(path):
    with path.open() as stream:
        return list(csv.DictReader(stream))


def make_quality_chart():
    rows = load_summary(ROOT / "results" / "final_benchmark_summary.csv")
    wanted = {}
    for row in rows:
        if int(row["n_qubits"]) == 2 and int(row["shots_per_setting"]) == 500:
            wanted[(row["state_type"], row["method"])] = float(row["mean_hs_distance"])
    methods = [("projected", "PSD projection", "4F718C"), ("low_rank", "Rank-1", "D59B35"), ("mle", "MLE", "0F7173")]
    states = [("haar", "Haar pure"), ("product", "Product pure"), ("mixed", "Full-rank mixed")]
    image = Image.new("RGB", (1500, 820), "white")
    draw = ImageDraw.Draw(image)
    title_font, label_font, small_font = find_font(42, True), find_font(28), find_font(23)
    draw.text((70, 40), "Two-qubit denoising error at 500 shots per setting", font=title_font, fill=(20, 52, 78))
    x0, y0, chart_w, chart_h = 140, 155, 1260, 510
    max_value = 0.56
    for i in range(6):
        x = x0 + chart_w * i / 5
        draw.line((x, y0, x, y0 + chart_h), fill=(220, 226, 231), width=2)
        draw.text((x - 18, y0 + chart_h + 18), f"{max_value*i/5:.2f}", font=small_font, fill=(70, 80, 90))
    bar_h = 42
    for si, (state, state_label) in enumerate(states):
        group_y = y0 + 25 + si * 160
        draw.text((10, group_y + 36), state_label, font=small_font, fill=(32, 38, 45))
        for mi, (method, method_label, color) in enumerate(methods):
            value = wanted[(state, method)]
            y = group_y + mi * 46
            width = chart_w * value / max_value
            draw.rounded_rectangle((x0, y, x0 + width, y + bar_h - 6), radius=8, fill="#" + color)
            draw.text((x0 + width + 12, y + 3), f"{value:.3f}", font=small_font, fill=(45, 50, 55))
    legend_y = 735
    for i, (_, label, color) in enumerate(methods):
        x = 285 + i * 340
        draw.rounded_rectangle((x, legend_y, x + 34, legend_y + 25), radius=5, fill="#" + color)
        draw.text((x + 48, legend_y - 4), label, font=small_font, fill=(45, 50, 55))
    path = FIG_DIR / "quality_hs.png"
    image.save(path)
    return path


def make_runtime_chart():
    rows = load_summary(ROOT / "results" / "final_benchmark_summary.csv")
    grouped = defaultdict(list)
    for row in rows:
        if int(row["n_qubits"]) == 2:
            grouped[row["method"]].append(float(row["mean_reconstruction_seconds"]))
    methods = [("linear", "Linear"), ("projected", "PSD"), ("low_rank", "Rank-1"), ("shrinkage", "Shrink"), ("mle", "MLE")]
    values = [sum(grouped[key]) / len(grouped[key]) * 1000 for key, _ in methods]
    image = Image.new("RGB", (1500, 760), "white")
    draw = ImageDraw.Draw(image)
    title_font, label_font, small_font = find_font(42, True), find_font(28), find_font(23)
    draw.text((70, 40), "Mean two-qubit NumPy reconstruction time", font=title_font, fill=(20, 52, 78))
    x0, y0, chart_w, chart_h = 150, 150, 1240, 440
    log_min, log_max = -1.1, 2.0
    for exponent in [-1, 0, 1, 2]:
        y = y0 + chart_h * (log_max - exponent) / (log_max - log_min)
        draw.line((x0, y, x0 + chart_w, y), fill=(220, 226, 231), width=2)
        draw.text((55, y - 14), f"10^{exponent} ms", font=small_font, fill=(70, 80, 90))
    slot = chart_w / len(methods)
    for i, ((_, label), value) in enumerate(zip(methods, values)):
        exponent = math.log10(value)
        top = y0 + chart_h * (log_max - exponent) / (log_max - log_min)
        left = x0 + i * slot + 45
        right = x0 + (i + 1) * slot - 45
        draw.rounded_rectangle((left, top, right, y0 + chart_h), radius=10, fill=(15, 113, 115))
        draw.text((left, top - 36), f"{value:.3f}", font=small_font, fill=(45, 50, 55))
        draw.text((left + 5, y0 + chart_h + 18), label, font=label_font, fill=(45, 50, 55))
    draw.text((570, 685), "Logarithmic vertical scale", font=small_font, fill=(93, 105, 116))
    path = FIG_DIR / "runtime.png"
    image.save(path)
    return path


def clean_inline(text):
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    return text.replace("^dagger", "^dagger")


def table_widths(headers):
    count = len(headers)
    if count == 2:
        return [2600, 6760]
    if count == 4:
        return [1900, 1500, 3200, 2760]
    if count == 5:
        return [1780, 1350, 2750, 1740, 1740]
    if count == 6:
        return [1550, 1020, 1220, 1800, 1870, 1900]
    base = 9360 // count
    widths = [base] * count
    widths[-1] += 9360 - sum(widths)
    return widths


def add_markdown_table(doc, rows):
    headers = rows[0]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = clean_inline(value)
        set_cell_shading(cell, LIGHT)
        for run in cell.paragraphs[0].runs:
            set_font(run, size=9, bold=True, color=NAVY)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    for row_index, values in enumerate(rows[1:]):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = clean_inline(value)
            if row_index % 2:
                set_cell_shading(cells[i], PALE)
            for run in cells[i].paragraphs[0].runs:
                set_font(run, size=8.5, color=INK)
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(0)
            cells[i].paragraphs[0].paragraph_format.line_spacing = 1.08
    set_table_geometry(table, table_widths(headers))
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_cover(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(78)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = kicker.add_run("NBQSS 2026 | TECHNICAL REPORT")
    set_font(r, size=11, bold=True, color=GOLD)
    kicker.paragraph_format.space_after = Pt(18)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Building a Hardware-Agnostic\nQuantum State Tomography\nand Denoising Suite")
    set_font(r, size=28, bold=True, color=NAVY)
    title.paragraph_format.space_after = Pt(18)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Methods, implementation, benchmarks, and a scalable roadmap beyond attention")
    set_font(r, size=14, color=TEAL)
    subtitle.paragraph_format.space_after = Pt(62)
    callout = doc.add_table(rows=1, cols=1)
    callout.style = "Table Grid"
    callout.cell(0, 0).text = (
        "Core recommendation\nStart with physical likelihood and validated structure. "
        "Use attention only when repeatable device noise justifies learned global corrections."
    )
    set_cell_shading(callout.cell(0, 0), LIGHT)
    set_table_geometry(callout, [9360])
    for i, run in enumerate(callout.cell(0, 0).paragraphs[0].runs):
        set_font(run, size=11, bold=i == 0, color=NAVY if i == 0 else INK)
    callout.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    callout.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(0)
    footer_space = doc.add_paragraph()
    footer_space.paragraph_format.space_after = Pt(42)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Version 1.0  |  17 August 2026\nReference implementation: nbqst 0.1.0")
    set_font(r, size=10, color=MUTED)
    doc.add_page_break()


def add_contents(doc):
    heading = doc.add_paragraph("Contents", style="Heading 1")
    entries = [
        "Executive recommendations",
        "Challenge interpretation and requirement mapping",
        "Attention-paper assessment and notebook audit",
        "Measurement, reconstruction, and denoising methods",
        "Hardware-agnostic architecture",
        "Experimental protocol and results",
        "Scaling roadmap, validation, and reproducibility",
        "Generative AI disclosure and references",
    ]
    for entry in entries:
        doc.add_paragraph(entry, style="List Bullet")
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    run = note.add_run("Reader's guide: Sections 1 and 6 contain the decision; Sections 4 and 9 contain the evidence; Sections 7 and 11 define the software roadmap.")
    set_font(run, size=10, italic=True, color=MUTED)
    doc.add_page_break()


def add_paragraph_with_inline(doc, text, style=None):
    paragraph = doc.add_paragraph(style=style)
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, name="Menlo", size=9.5, color=NAVY)
        elif part.startswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, bold=True, color=INK)
        else:
            paragraph.add_run(part)
    return paragraph


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    quality_chart = make_quality_chart()
    runtime_chart = make_runtime_chart()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("NBQSS 2026  |  Hardware-Agnostic QST")
    set_font(r, size=8.5, color=MUTED)
    add_page_number(section.footer.paragraphs[0])
    add_cover(doc)
    add_contents(doc)

    lines = REPORT_MD.read_text(encoding="utf-8").splitlines()
    index = 6  # skip title, subtitle, version line, and surrounding blanks
    paragraph_buffer = []
    in_code = False
    code_lines = []
    current_num_id = None

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            add_paragraph_with_inline(doc, " ".join(s.strip() for s in paragraph_buffer))
            paragraph_buffer = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped.startswith("```"):
            current_num_id = None
            flush_paragraph()
            if in_code:
                p = doc.add_paragraph(style="Code Block")
                p.add_run("\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not stripped:
            flush_paragraph()
            current_num_id = None
            index += 1
            continue
        if stripped.startswith("|"):
            flush_paragraph()
            current_num_id = None
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
                index += 1
            rows = [table_lines[0]] + table_lines[2:]
            add_markdown_table(doc, rows)
            continue
        if stripped.startswith("### "):
            flush_paragraph()
            current_num_id = None
            title = clean_inline(stripped[4:])
            doc.add_paragraph(title, style="Heading 2")
            if title.startswith("9.1"):
                doc.add_picture(str(quality_chart), width=Inches(6.35))
                doc.add_paragraph(
                    "Figure 1. Mean Hilbert-Schmidt error for physical two-qubit estimators at 500 shots per setting. Rank one is appropriate for pure targets and misspecified for mixed targets.",
                    style="Caption",
                )
            if title.startswith("9.3"):
                doc.add_picture(str(runtime_chart), width=Inches(6.35))
                doc.add_paragraph(
                    "Figure 2. Eager NumPy reconstruction time averaged across two-qubit benchmark conditions. The vertical scale is logarithmic.",
                    style="Caption",
                )
            index += 1
            continue
        if stripped.startswith("## "):
            flush_paragraph()
            current_num_id = None
            doc.add_paragraph(clean_inline(stripped[3:]), style="Heading 1")
            index += 1
            continue
        if re.match(r"^\d+\. ", stripped):
            flush_paragraph()
            if current_num_id is None:
                current_num_id = create_decimal_numbering(doc)
            paragraph = add_paragraph_with_inline(doc, re.sub(r"^\d+\. ", "", stripped), style="List Number")
            apply_numbering(paragraph, current_num_id)
            index += 1
            continue
        if stripped.startswith("- "):
            flush_paragraph()
            current_num_id = None
            add_paragraph_with_inline(doc, stripped[2:], style="List Bullet")
            index += 1
            continue
        paragraph_buffer.append(stripped)
        current_num_id = None
        index += 1
    flush_paragraph()

    core = doc.core_properties
    core.title = "Building a Hardware-Agnostic Quantum State Tomography and Denoising Suite"
    core.subject = "NBQSS 2026 tomography challenge technical report"
    core.author = "NBQST project team"
    core.keywords = "quantum state tomography, denoising, Array API, MLE, low rank"
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
